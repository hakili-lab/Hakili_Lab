"""
TestRegistry — Catalogue des tests standards Hakili Lab.

Pour chaque test pré-défini, charge automatiquement :
  - L'énoncé (texte extrait du fichier DOCX dans data/Documents/)
  - Le barème Rubric (construit depuis le YAML data/knowledge/bareme_test_*.yaml)
  - Le bareme_id RAG correspondant

Utilisation dans le pipeline :
    registry = get_registry()
    test = registry.get_test("hakili_3e_v1")
    rubric      = test.rubric        # Rubric Pydantic prêt à l'emploi
    subject_text = test.subject_text  # Texte de l'énoncé
    bareme_id    = test.bareme_id

Tests archivés
--------------
Les 6 tests construits avant le référentiel Urie v2 sont marqués `archive: True`
(arbitrage A du 2026-07-30, voir docs/harmonisation_donnees.md) : leurs sujets ont
été refaits et ils ne doivent plus être proposés pour une nouvelle correction.

Ils restent chargés, et c'est délibéré : `get_test()` doit continuer à les résoudre
pour les copies déjà corrigées — le pipeline s'en sert pour retrouver les niveaux
déclarés d'un test lors de la détermination de la classe (`_apply_extracted_classe`).
Les retirer du catalogue casserait la relecture de l'historique.

    available_tests()  -> tests proposables pour une NOUVELLE correction
    all_tests()        -> tous les tests, archivés compris (historique)
    get_test(id)       -> résout n'importe quel test, archivé ou non
"""
from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import TYPE_CHECKING

import yaml

from src.knowledge.answer_loader import get_answer_loader

if TYPE_CHECKING:
    from src.models.domain import Rubric

logger = logging.getLogger(__name__)

_ROOT = Path(__file__).parent.parent.parent
_DOCS_DIR = _ROOT / "data" / "Documents"
_KB_DIR = _ROOT / "data" / "knowledge"

# ── Catalogue des tests Hakili ────────────────────────────────────────────────
#
# `archive: True` = test d'avant le référentiel Urie v2, remplacé par les nouveaux
# sujets à cadres ancrés. Chargé mais plus proposé — voir le docstring du module.

_TEST_CATALOG: dict[str, dict] = {
    # ── Tests Urie v2 — actifs ────────────────────────────────────────────────
    # Barèmes générés depuis Referentiel_Urie_v0.xlsx par
    # scripts/generer_baremes_urie.py. `label` et `niveaux` sont laissés vides :
    # ils sont lus dans le `meta` du YAML (titre, classe) pour éviter de dupliquer
    # la même information à deux endroits qui divergeraient.
    #
    # Pas de DOCX : les sujets sont des PDF dont les mathématiques sont dessinées
    # en vectoriel, donc non extractibles en texte. Ce n'est pas bloquant — dans
    # ce format l'élève compose SUR le sujet, donc la copie scannée porte déjà
    # l'énoncé imprimé que l'IA transcrit.
    #
    # Pas de corrigé complet : le classeur ne donne la bonne réponse que pour les
    # 71 QCM (sur 280). Voir docs/harmonisation_donnees.md, arbitrage B.
    **{
        f"urie_{niveau}": {
            "label": "",
            "description": (
                "40 questions · partie A (30 courtes) + partie B (10 rédigées) · "
                "corrigé disponible pour les QCM uniquement"
            ),
            "niveaux": "",
            "docx_filename": "",
            "bareme_yaml": f"bareme_urie_{niveau}.yaml",
            "corrige_yaml": "",
            # Le sujet PDF sert à l'IMPRESSION, pas à l'extraction de texte : ses
            # mathématiques sont vectorielles (voir docs/harmonisation_donnees.md
            # §3bis). L'enseignant doit pouvoir l'imprimer depuis l'application.
            "sujet_pdf": f"Test_diagnostique_entree_{niveau}.pdf",
        }
        for niveau in ("6eme", "5eme", "4eme", "3eme", "2ndeC", "1ereD", "tleD")
    },
    # ── Tests d'avant le référentiel — archivés (D-CEO-26) ────────────────────
    "hakili_3e_v1": {
        "label": "Test diagnostique Hakili 3e v1",
        "description": "Évalue les acquis de la 6e à la 4e (calcul et géométrie)",
        "niveaux": "6e · 5e · 4e",
        "docx_filename": "Hakilisso test de niveau 3e.docx",
        "bareme_yaml": "bareme_test_3e.yaml",
        "corrige_yaml": "corrige_test_3e.yaml",
        "archive": True,
    },
    "hakili_3e_v2": {
        "label": "Test diagnostique Hakili 3e v2",
        "description": "Évalue les acquis de la 6e à la 4e (calcul et géométrie) — version 2",
        "niveaux": "6e · 5e · 4e",
        "docx_filename": "Test-niveau 3ieme v2.docx",
        "bareme_yaml": "bareme_test_3e_v2.yaml",
        "corrige_yaml": "corrige_test_3e_v2.yaml",
        "archive": True,
    },
    "hakili_6e_v1": {
        "label": "Test diagnostique Hakili 6e v1",
        "description": "Évalue les acquis du CE1 au CM2 (calcul et géométrie)",
        "niveaux": "CE1 · CE2 · CM1 · CM2",
        "docx_filename": "TEST DE NIVEAU,6eme,GROUPE 1.docx",
        "bareme_yaml": "bareme_test_6e.yaml",
        "corrige_yaml": "corrige_test_6e.yaml",
        "archive": True,
    },
    "hakili_2ndeC_v1": {
        "label": "Test diagnostique Hakili 2nde C v1",
        "description": "Évaluation diagnostique niveau 2nde C (numériques et géométriques)",
        "niveaux": "6e · 5e · 4e · 3e · 2nde",
        "docx_filename": "EVALUATION 2ndC.docx",
        "bareme_yaml": "bareme_test_2ndeC.yaml",
        "corrige_yaml": "corrige_test_2ndeC.yaml",
        "archive": True,
    },
    "hakili_4e_v1": {
        "label": "Évaluation de Mathématiques 4ème v1",
        "description": "Évalue les compétences 6e → 4e (ensembles, calculs, équations, géométrie)",
        "niveaux": "6e · 5e · 4e",
        "docx_filename": "",   # Pas de DOCX — seul le corrigé PDF est disponible
        "bareme_yaml": "bareme_test_4e.yaml",
        "corrige_yaml": "corrige_test_4e.yaml",
        "archive": True,
    },
    "hakili_tle_v1": {
        "label": "Devoir de Mathématiques Terminale v1",
        "description": "Évalue les compétences 4e → Tle (puissances, racines, factorisation, vecteurs, trigo)",
        "niveaux": "4e · 3e · Tle",
        "docx_filename": "",   # Pas de DOCX — seul le corrigé PDF est disponible
        "bareme_yaml": "bareme_test_tle.yaml",
        "corrige_yaml": "corrige_test_tle.yaml",
        "archive": True,
    },
}


@dataclass
class HakiliTest:
    """Test Hakili pré-chargé — prêt à injecter dans run_single_copy()."""
    bareme_id: str
    label: str
    description: str
    niveaux: str
    subject_text: str
    rubric: "Rubric"  # type: ignore[type-arg]
    total_questions: int
    official_answers: str = ""  # Corrigé officiel formaté pour injection dans le prompt
    archive: bool = False       # True = plus proposé pour une nouvelle correction
    sujet_pdf: Path | None = None  # sujet imprimable, None si absent
    # {code question: format} — vide pour les tests archivés, qui n'ont pas de
    # cadres ancrés. Sert au garde-fou de la lecture par zones (module 2).
    formats: dict[str, str] = field(default_factory=dict)


def _extract_docx_text(docx_path: Path) -> str:
    """Extrait le texte brut d'un fichier DOCX (tous paragraphes non vides)."""
    try:
        import docx  # python-docx
        doc = docx.Document(str(docx_path))
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
        # Inclure aussi les tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in lines:
                        lines.append(text)
        return "\n".join(lines)
    except ImportError:
        logger.error("python-docx non installé — impossible d'extraire l'énoncé DOCX.")
        return ""
    except Exception as exc:
        logger.error("Erreur extraction DOCX %s : %s", docx_path.name, exc)
        return ""


def _build_rubric_from_yaml(bareme_yaml_path: Path) -> tuple["Rubric", int, dict]:
    """
    Construit un objet Rubric Pydantic à partir d'un fichier de barème YAML.

    Deux formats sont acceptés :

    - **Urie v2** (`bareme_urie_*.yaml`, généré par `scripts/generer_baremes_urie.py`) :
      liste plate `questions`, clés `code` / `objet` / `bareme`. Structure plate parce
      que les 7 domaines du référentiel (N, L, G, D, F, M, S, T) ne rentrent pas dans
      un découpage numérique/géométrique.
    - **Ancien** (`bareme_test_*.yaml`, tests archivés) : sections
      `questions_numeriques` / `questions_geometriques`, clés `id` / `enonce_court` /
      `points_originaux`. Conservé pour la relecture de l'historique.

    Retourne (rubric, total_questions, meta).
    """
    from src.models.domain import Rubric, RubricItem

    try:
        raw = yaml.safe_load(bareme_yaml_path.read_text(encoding="utf-8"))
    except Exception as exc:
        logger.error("Erreur lecture barème YAML %s : %s", bareme_yaml_path.name, exc)
        return Rubric(subject="mathematics", total_points=0, items=[]), 0, {}

    meta = raw.get("meta", {}) or {}
    items: list[RubricItem] = []

    if "questions" in raw:  # format Urie v2
        for q in raw.get("questions", []):
            qid = q.get("code", "")
            if qid:
                items.append(
                    RubricItem(
                        id=qid,
                        label=q.get("objet", "") or qid,
                        max_score=float(q.get("bareme", 0.0)),
                    )
                )
    else:  # format ancien
        for section in ("questions_numeriques", "questions_geometriques"):
            for q in raw.get(section, []):
                qid = q.get("id", "")
                label = q.get("enonce_court", qid)
                pts = float(q.get("points_originaux", 1.0))
                if qid:
                    items.append(RubricItem(id=qid, label=label, max_score=pts))

    # Format de chaque question (qcm / court / redige / construction). Seuls les
    # barèmes Urie v2 le portent — il vient du classeur. `src/pipeline/zones.py`
    # le confronte à la géométrie des cadres du sujet : si les deux divergent,
    # c'est que le sujet scanné n'est pas de la version chargée en base, et on
    # l'apprend avant d'attribuer des réponses aux mauvaises questions.
    meta["formats"] = {
        q["code"]: q["format"]
        for q in (raw.get("questions") or [])
        if q.get("code") and q.get("format")
    }

    total_pts = float(meta.get("total_possible") or sum(i.max_score for i in items))
    rubric = Rubric(subject="mathematics", total_points=total_pts, items=items)
    return rubric, len(items), meta


class TestRegistry:
    """
    Singleton — chargé une fois en début de session.

    Charge et met en cache les tests Hakili (énoncés + barèmes).
    """

    def __init__(self) -> None:
        self._tests: dict[str, HakiliTest] = {}
        self._load_all()

    def _load_all(self) -> None:
        for test_id, cfg in _TEST_CATALOG.items():
            docx_filename = cfg.get("docx_filename", "")
            docx_path = (_DOCS_DIR / docx_filename) if docx_filename else None
            yaml_path = _KB_DIR / cfg["bareme_yaml"]

            archive = bool(cfg.get("archive", False))
            if docx_path is None:
                logger.info("Test '%s' — pas de DOCX (énoncé absent).", test_id)
            elif not docx_path.exists():
                # Pour un test archivé, l'absence du DOCX est normale : les sujets
                # ont été refaits et les anciens fichiers retirés. Le signaler comme
                # un avertissement laisserait croire à une installation cassée.
                niveau = logger.info if archive else logger.warning
                niveau(
                    "Test '%s' — DOCX absent%s : %s",
                    test_id,
                    " (archivé, attendu)" if archive else "",
                    docx_path.name,
                )
            if not yaml_path.exists():
                logger.warning(
                    "Test '%s' — barème YAML introuvable : %s", test_id, yaml_path
                )
                continue

            subject_text = (
                _extract_docx_text(docx_path)
                if docx_path is not None and docx_path.exists()
                else ""
            )
            sujet_pdf_nom = cfg.get("sujet_pdf", "")
            sujet_pdf = (_DOCS_DIR / sujet_pdf_nom) if sujet_pdf_nom else None
            if sujet_pdf is not None and not sujet_pdf.exists():
                logger.warning(
                    "Test '%s' — sujet PDF introuvable : %s (impression indisponible)",
                    test_id, sujet_pdf.name,
                )
                sujet_pdf = None

            rubric, total, bareme_meta = _build_rubric_from_yaml(yaml_path)

            # Chargement du corrigé officiel
            official_answers = get_answer_loader().get_official_answers(test_id)

            # Les tests Urie v2 laissent label/niveaux vides dans le catalogue :
            # la valeur fait autorité dans le `meta` du barème, pas ici.
            label = cfg["label"] or str(bareme_meta.get("titre") or test_id)
            niveaux = cfg["niveaux"] or str(bareme_meta.get("classe") or "")

            self._tests[test_id] = HakiliTest(
                bareme_id=test_id,
                label=label,
                description=cfg["description"],
                niveaux=niveaux,
                subject_text=subject_text,
                rubric=rubric,
                total_questions=total,
                official_answers=official_answers,
                archive=archive,
                sujet_pdf=sujet_pdf,
                formats=bareme_meta.get("formats") or {},
            )
            corrige_status = f"corrigé {len(official_answers)} chars" if official_answers else "sans corrigé"
            logger.info(
                "Test chargé : %s (%d questions, énoncé %d chars, %s)%s",
                test_id, total, len(subject_text), corrige_status,
                " [ARCHIVÉ]" if cfg.get("archive") else "",
            )

    def available_tests(self) -> dict[str, HakiliTest]:
        """Tests proposables pour une NOUVELLE correction — sans les archivés."""
        return {tid: t for tid, t in self._tests.items() if not t.archive}

    def all_tests(self) -> dict[str, HakiliTest]:
        """Tous les tests chargés, archivés compris — relecture de l'historique."""
        return dict(self._tests)

    def get_test(self, test_id: str) -> HakiliTest | None:
        """Résout n'importe quel test, archivé ou non — l'historique en dépend."""
        return self._tests.get(test_id)

    @property
    def ids(self) -> list[str]:
        """Identifiants de tous les tests chargés, archivés compris."""
        return list(self._tests.keys())


# ── Singleton module-level ────────────────────────────────────────────────────

_registry: TestRegistry | None = None


def get_registry() -> TestRegistry:
    global _registry
    if _registry is None:
        _registry = TestRegistry()
    return _registry
